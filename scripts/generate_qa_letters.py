"""Generate and audit one non-production DOCX for every active letter type.

The script exercises the real Flask ``/generate`` route through ``create_app``
and its test client. Master records are synthetic test fixtures and SQLite state
lives in a temporary directory. Generated artifacts are written to
``qa/generated`` (ignored by Git).

Run from the repository root:
    python scripts/generate_qa_letters.py
"""

from __future__ import annotations

import hashlib
import io
import json
import sqlite3
import sys
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path

import docx


ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT / "qa" / "generated"
FIXTURE_DATA_DIR = ROOT / "tests" / "fixtures" / "master"
PROTECTED_DATABASES = (
    ROOT / "data" / "runtime" / "surat_smada.db",
    ROOT / "data" / "runtime" / "data_surat.db",
)
UNRESOLVED_MARKERS = ("{{", "{%", "{#", "${")

EXPECTED_TABLE_ROWS = {
    "izin_guru": 8,
    "cuti_guru": 10,
    "sakit_guru": 8,
    "surat_tugas_guru": 7,
    "surat_keterangan_guru": 6,
    "izin_murid": 7,
    "dispensasi_murid": 9,
}

CASE_FIELDS = {
    "izin_guru": {
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Keperluan keluarga & administrasi",
    },
    "cuti_guru": {
        "jenis_cuti": "Cuti Tahunan",
        "lama_cuti": "2 hari",
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Keperluan keluarga & administrasi",
        "alamat_selama_cuti": "Wonosari, Gunungkidul",
    },
    "sakit_guru": {
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Pemulihan kesehatan & pemeriksaan dokter",
    },
    "surat_tugas_guru": {
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Rapat koordinasi & evaluasi pendidikan",
    },
    "surat_keterangan_guru": {
        "tanggal_mulai": "2026-08-24",
        "keperluan": "Masih aktif melaksanakan tugas di SMA Negeri 2 Wonosari",
    },
    "izin_murid": {
        "nama_wali": "WALI QA E-SURAT",
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Keperluan keluarga & pendampingan orang tua",
    },
    "dispensasi_murid": {
        "nama_kegiatan": "Festival Pelajar dan Seni",
        "penyelenggara": "Dinas Pendidikan & Kebudayaan",
        "tempat_kegiatan": "Wonosari",
        "tanggal_mulai": "2026-08-24",
        "tanggal_selesai": "2026-08-25",
        "keperluan": "Mengikuti kegiatan sebagai perwakilan sekolah",
    },
}


def _file_fingerprint(path: Path) -> tuple[int, int, str] | None:
    if not path.exists():
        return None
    stat = path.stat()
    return stat.st_size, stat.st_mtime_ns, hashlib.sha256(path.read_bytes()).hexdigest()


def _assert_valid_docx(payload: bytes, jenis: str, expected_signer: dict[str, str], subject_name: str) -> None:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        broken = archive.testzip()
        if broken:
            raise AssertionError(f"{jenis}: entry DOCX rusak: {broken}")
        word_xml = "\n".join(
            archive.read(name).decode("utf-8", errors="replace")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
        unresolved = [marker for marker in UNRESOLVED_MARKERS if marker in word_xml]
        if unresolved:
            raise AssertionError(f"{jenis}: token belum terisi: {unresolved}")

    document = docx.Document(io.BytesIO(payload))
    if len(document.tables) != 1:
        raise AssertionError(f"{jenis}: tabel={len(document.tables)}, seharusnya 1")
    if len(document.tables[0].rows) != EXPECTED_TABLE_ROWS[jenis]:
        raise AssertionError(
            f"{jenis}: baris tabel={len(document.tables[0].rows)}, "
            f"seharusnya {EXPECTED_TABLE_ROWS[jenis]}"
        )

    top_level = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    try:
        signature_start = next(index for index, text in enumerate(top_level) if text.startswith("Wonosari,"))
    except StopIteration as exc:
        raise AssertionError(f"{jenis}: blok tanda tangan tidak ditemukan") from exc
    signature = top_level[signature_start:]

    for expected in (expected_signer["peran"], expected_signer["nama"]):
        if expected not in signature:
            raise AssertionError(f"{jenis}: penandatangan tidak tepat; {expected!r} tidak ditemukan")

    signer_nip = expected_signer.get("nip", "")
    expected_id_line = f"NIP. {signer_nip}" if signer_nip else ""
    if expected_id_line and expected_id_line not in signature:
        raise AssertionError(f"{jenis}: NIP penandatangan tidak ditemukan")
    if not expected_id_line and any(text.startswith("NIP.") for text in signature):
        raise AssertionError(f"{jenis}: baris NIP seharusnya tidak ditampilkan")
    if subject_name != expected_signer["nama"] and subject_name in signature:
        raise AssertionError(f"{jenis}: subjek surat keliru menjadi penandatangan")


def _expected_signer(jenis: str, info: dict, person: dict, state: dict, fields: dict) -> dict[str, str]:
    signer_kind = str(info["signer"])
    if signer_kind == "pemohon":
        return {"nama": person["nama"], "nip": person.get("nip", ""), "peran": "Pemohon"}
    if signer_kind == "wali":
        return {"nama": fields["nama_wali"], "nip": "", "peran": "Orang Tua / Wali"}
    kepsek = state["kepsek"]
    return {"nama": kepsek["nama"], "nip": kepsek["nip"], "peran": "Kepala Sekolah"}


def main() -> None:
    protected_before = {path: _file_fingerprint(path) for path in PROTECTED_DATABASES}
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if str(ROOT) not in sys.path:
        sys.path.insert(0, str(ROOT))

    # Import dilakukan sesudah fingerprint database produksi dicatat. Global app
    # kompatibilitas tidak menginisialisasi DB karena INIT_DB_ON_CREATE=False.
    from esurat import JENIS_SURAT, WIB, create_app

    manifest: list[dict[str, str | int]] = []
    try:
        with tempfile.TemporaryDirectory(prefix="esurat-qa-db-") as temp_dir:
            qa_database = Path(temp_dir) / "qa.sqlite3"
            qa_app = create_app(
                {
                    "TESTING": True,
                    "DATABASE": qa_database,
                    "DATA_DIR": FIXTURE_DATA_DIR,
                    "INIT_DB_ON_CREATE": True,
                    "AUTH_ENABLED": False,
                    "BIND_HOST": "127.0.0.1",
                    "SECRET_KEY": "qa-only-not-for-production",
                    "NOW_FUNC": lambda: datetime(2026, 8, 23, 10, 30, tzinfo=WIB),
                }
            )
            state = qa_app.extensions["esurat_data"]
            kepsek_nip = state["kepsek"]["nip"]
            guru = next(record for record in state["guru"] if record["nip"] != kepsek_nip)
            murid = state["murid"][0]

            with qa_app.test_client() as client:
                csrf_response = client.get("/api/csrf")
                if csrf_response.status_code != 200:
                    raise AssertionError(f"CSRF bootstrap gagal: HTTP {csrf_response.status_code}")
                csrf_token = csrf_response.get_json()["csrf_token"]

                for jenis, info in JENIS_SURAT.items():
                    person = guru if info["kategori"] == "guru" else murid
                    fields = dict(CASE_FIELDS[jenis])
                    request_id = f"qa-integration-{jenis.replace('_', '-')}"
                    form = {
                        "_csrf_token": csrf_token,
                        "request_id": request_id,
                        "jenis_surat": jenis,
                        "kategori": info["kategori"],
                        "id_value": person.get("nip") or person["nis"],
                        "tanggal_surat": "2026-08-23",
                        "kode_arsip": info["default_kode"],
                        **fields,
                    }
                    response = client.post("/generate", data=form)
                    if response.status_code != 200:
                        detail = response.get_json(silent=True) or response.get_data(as_text=True)
                        raise AssertionError(f"{jenis}: HTTP {response.status_code}: {detail}")

                    expected_signer = _expected_signer(jenis, info, person, state, fields)
                    _assert_valid_docx(response.data, jenis, expected_signer, person["nama"])
                    output_path = OUTPUT_DIR / f"{jenis}.docx"
                    output_path.write_bytes(response.data)
                    manifest.append(
                        {
                            "jenis": jenis,
                            "output": output_path.relative_to(ROOT).as_posix(),
                            "bytes": len(response.data),
                            "sha256": hashlib.sha256(response.data).hexdigest(),
                            "nomor_surat": response.headers.get("X-Letter-Number", ""),
                            "request_id": response.headers.get("X-Request-ID", ""),
                            "penandatangan": expected_signer["nama"],
                            "peran": expected_signer["peran"],
                        }
                    )
                    print(
                        f"[QA OK] {jenis}: signer={expected_signer['peran']} / "
                        f"{expected_signer['nama']}"
                    )

            connection = sqlite3.connect(qa_database)
            try:
                statuses = connection.execute(
                    "SELECT status, COUNT(*) FROM riwayat_surat GROUP BY status ORDER BY status"
                ).fetchall()
            finally:
                connection.close()
            if statuses != [("generated", len(JENIS_SURAT))]:
                raise AssertionError(f"Status database QA tidak sesuai: {statuses}")

        manifest_path = OUTPUT_DIR / "manifest.json"
        manifest_path.write_text(
            json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        print(f"[QA OK] {len(manifest)} DOCX tersimpan di {OUTPUT_DIR}")
        print(f"[QA OK] Manifest: {manifest_path}")
    finally:
        protected_after = {path: _file_fingerprint(path) for path in PROTECTED_DATABASES}
        changed = [str(path) for path in PROTECTED_DATABASES if protected_before[path] != protected_after[path]]
        if changed:
            raise AssertionError(f"Database produksi berubah selama QA: {', '.join(changed)}")
        print("[QA OK] Database produksi tidak berubah.")


if __name__ == "__main__":
    main()
