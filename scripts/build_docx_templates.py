"""
Bangun ulang tujuh template DOCX aktif E-Surat SMADA.

Builder ini sengaja memakai master immutable yang tidak pernah menjadi target
output. Tiga elemen awal body master dipertahankan sebagai blok kop; seluruh
elemen sesudahnya (paragraf maupun tabel) dibuang sebelum isi surat dibuat.
Dengan demikian, menjalankan script berulang kali tidak menggandakan isi.

Cara menjalankan:
    python scripts/build_docx_templates.py
"""

from __future__ import annotations

import hashlib
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

import docx
from docxtpl import DocxTemplate
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_ROOT = BASE_DIR / "templates_surat"
ACTIVE_TEMPLATE_DIR = TEMPLATE_ROOT / "active"

# Master ini hanya sumber kop. Jangan pernah masukkan namanya ke TEMPLATE_SPECS.
MASTER_DOC = TEMPLATE_ROOT / "master" / "kop_smada.docx"
KOP_BODY_ELEMENT_COUNT = 3

FONT_NAME = "Times New Roman"
BODY_FONT_SIZE_PT = 12
TITLE_FONT_SIZE_PT = 14

# A4 21 cm, margin kiri 3 cm dan kanan 2 cm menghasilkan lebar isi 16 cm.
TABLE_COLUMN_WIDTHS_DXA = (2948, 227, 5896)
TABLE_WIDTH_DXA = sum(TABLE_COLUMN_WIDTHS_DXA)
CELL_MARGIN_DXA = {"top": 80, "left": 100, "bottom": 80, "right": 100}


@dataclass(frozen=True)
class TemplateSpec:
    filename: str
    title: str
    opening: str
    rows: tuple[tuple[str, str], ...]
    closing: str
    multi_students: bool = False


TEMPLATE_SPECS: tuple[TemplateSpec, ...] = (
    TemplateSpec(
        filename="izin_guru.docx",
        title="SURAT PERMOHONAN IZIN PEGAWAI",
        opening="Yang bertanda tangan di bawah ini mengajukan permohonan izin tidak masuk kerja:",
        rows=(
            ("Nama", "{{ nama }}"),
            ("NIP", "{{ nip }}"),
            ("Jabatan", "{{ jabatan }}"),
            ("Pangkat / Golongan", "{{ golongan }}"),
            ("Unit Kerja", "{{ unit_kerja }}"),
            ("Tanggal Mulai", "{{ tanggal_mulai }}"),
            ("Tanggal Selesai", "{{ tanggal_selesai }}"),
            ("Keperluan / Alasan", "{{ keperluan }}"),
        ),
        closing="Demikian permohonan izin ini dibuat untuk dapat dipergunakan sebagaimana mestinya.",
    ),
    TemplateSpec(
        filename="cuti_guru.docx",
        title="SURAT PERMOHONAN CUTI PEGAWAI",
        opening="Yang bertanda tangan di bawah ini mengajukan permohonan cuti pegawai:",
        rows=(
            ("Nama", "{{ nama }}"),
            ("NIP", "{{ nip }}"),
            ("Jabatan", "{{ jabatan }}"),
            ("Pangkat / Golongan", "{{ golongan }}"),
            ("Jenis Cuti", "{{ jenis_cuti }}"),
            ("Lama Cuti", "{{ lama_cuti }}"),
            ("Terhitung Mulai Tgl", "{{ tanggal_mulai }}"),
            ("Sampai Dengan Tgl", "{{ tanggal_selesai }}"),
            ("Alasan Cuti", "{{ keperluan }}"),
            ("Alamat Selama Cuti", "{{ alamat_selama_cuti }}"),
        ),
        closing="Demikian permohonan cuti ini disampaikan untuk pertimbangan lebih lanjut.",
    ),
    TemplateSpec(
        filename="sakit_guru.docx",
        title="SURAT PEMBERITAHUAN SAKIT",
        opening="Memberitahukan bahwa pegawai bersangkutan tidak dapat melaksanakan tugas karena sakit:",
        rows=(
            ("Nama", "{{ nama }}"),
            ("NIP", "{{ nip }}"),
            ("Jabatan", "{{ jabatan }}"),
            ("Pangkat / Golongan", "{{ golongan }}"),
            ("Unit Kerja", "{{ unit_kerja }}"),
            ("Tanggal Mulai Sakit", "{{ tanggal_mulai }}"),
            ("Sampai Tanggal", "{{ tanggal_selesai }}"),
            ("Keterangan Sakit", "{{ keperluan }}"),
        ),
        closing="Demikian surat pemberitahuan sakit ini dibuat dengan sebenarnya.",
    ),
    TemplateSpec(
        filename="3. Surat Tugas-smada.docx",
        title="SURAT TUGAS",
        opening="Kepala SMA Negeri 2 Wonosari dengan ini menugaskan kepada:",
        rows=(
            ("Nama", "{{ nama }}"),
            ("NIP", "{{ nip }}"),
            ("Pangkat / Golongan", "{{ golongan }}"),
            ("Jabatan", "{{ jabatan }}"),
            ("Tanggal Tugas", "{{ tanggal_mulai }}"),
            ("Sampai Tanggal", "{{ tanggal_selesai }}"),
            ("Uraian / Keperluan Tugas", "{{ keperluan }}"),
        ),
        closing="Demikian surat tugas ini diberikan untuk dilaksanakan dengan penuh rasa tanggung jawab.",
    ),
    TemplateSpec(
        filename="11. Surat Keterangan-smada.docx",
        title="SURAT KETERANGAN",
        opening="Kepala SMA Negeri 2 Wonosari dengan ini menerangkan bahwa:",
        rows=(
            ("Nama", "{{ nama }}"),
            ("NIP", "{{ nip }}"),
            ("Pangkat / Golongan", "{{ golongan }}"),
            ("Jabatan", "{{ jabatan }}"),
            ("Terhitung Mulai Tgl", "{{ tanggal_mulai }}"),
            ("Menerangkan Bahwa", "{{ keperluan }}"),
        ),
        closing="Demikian surat keterangan ini dibuat dengan sebenarnya untuk dapat dipergunakan sebagaimana mestinya.",
    ),
    TemplateSpec(
        filename="izin_murid.docx",
        title="SURAT IZIN SISWA",
        opening="Memberitahukan bahwa siswa bersangkutan mengajukan izin tidak mengikuti kegiatan belajar mengajar:",
        rows=(
            ("Nama Siswa", "{{ nama }}"),
            ("NIS / NISN", "{{ nis }} / {{ nisn }}"),
            ("Kelas", "Kelas {{ kelas }}"),
            ("Tanggal Mulai Izin", "{{ tanggal_mulai }}"),
            ("Tanggal Selesai Izin", "{{ tanggal_selesai }}"),
            ("Alasan Izin", "{{ keperluan }}"),
            ("Orang Tua / Wali", "{{ nama_wali }}"),
        ),
        closing="Demikian surat izin ini disampaikan agar menjadi maklum.",
    ),
    TemplateSpec(
        filename="dispensasi_murid.docx",
        title="SURAT DISPENSASI SISWA",
        opening="Kepala SMA Negeri 2 Wonosari memberikan dispensasi kepada siswa:",
        rows=(
            ("Nama Kegiatan", "{{ nama_kegiatan }}"),
            ("Penyelenggara", "{{ penyelenggara }}"),
            ("Tempat Kegiatan", "{{ tempat_kegiatan }}"),
            ("Tanggal Mulai", "{{ tanggal_mulai }}"),
            ("Tanggal Selesai", "{{ tanggal_selesai }}"),
            ("Uraian / Keperluan", "{{ keperluan }}"),
        ),
        closing="Demikian surat dispensasi ini dibuat untuk dipergunakan sebagaimana mestinya.",
        multi_students=True,
    ),
)


def _get_or_add(parent, tag: str, *, first: bool = False):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        if first:
            parent.insert(0, child)
        else:
            parent.append(child)
    return child


def _set_dxa(element, value: int) -> None:
    element.set(qn("w:w"), str(value))
    element.set(qn("w:type"), "dxa")


def _set_font_properties(r_pr, size_pt: int) -> None:
    r_fonts = _get_or_add(r_pr, "w:rFonts", first=True)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)

    half_points = str(size_pt * 2)
    _get_or_add(r_pr, "w:sz").set(qn("w:val"), half_points)
    _get_or_add(r_pr, "w:szCs").set(qn("w:val"), half_points)


def set_run_font(run, *, size_pt: int = BODY_FONT_SIZE_PT, bold=None, underline=None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline
    _set_font_properties(run._element.get_or_add_rPr(), size_pt)


def configure_default_font(doc: docx.Document) -> None:
    normal = next((style for style in doc.styles if style.style_id == "Normal"), None)
    if normal is None:
        raise RuntimeError("Style Normal tidak ditemukan pada master template.")
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_FONT_SIZE_PT)
    normal_r_pr = normal.element.get_or_add_rPr()
    _set_font_properties(normal_r_pr, BODY_FONT_SIZE_PT)

    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    r_pr_default = _get_or_add(doc_defaults, "w:rPrDefault", first=True)
    default_r_pr = _get_or_add(r_pr_default, "w:rPr")
    _set_font_properties(default_r_pr, BODY_FONT_SIZE_PT)


def configure_page(doc: docx.Document) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(1.5)
        section.right_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(1.0)


def retain_kop_only(doc: docx.Document) -> None:
    body = doc._element.body
    content = [child for child in body if child.tag != qn("w:sectPr")]
    if len(content) < KOP_BODY_ELEMENT_COUNT:
        raise RuntimeError("Master tidak memiliki blok kop yang diharapkan.")

    kop = content[:KOP_BODY_ELEMENT_COUNT]
    drawing_count = sum(1 for element in kop for node in element.iter() if node.tag == qn("w:drawing"))
    if drawing_count != 1:
        raise RuntimeError(f"Blok kop master harus memuat tepat satu drawing; ditemukan {drawing_count}.")

    for element in content[KOP_BODY_ELEMENT_COUNT:]:
        body.remove(element)


def add_text_paragraph(
    doc: docx.Document,
    text: str,
    *,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    size_pt: int = BODY_FONT_SIZE_PT,
    bold: bool = False,
    underline: bool = False,
    space_before_pt: int = 0,
    space_after_pt: int = 0,
    line_spacing: float = 1.15,
    keep_with_next: bool = False,
    left_indent_cm: float | None = None,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if left_indent_cm is not None:
        paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    run = paragraph.add_run(text)
    set_run_font(run, size_pt=size_pt, bold=bold, underline=underline)
    return paragraph


def remove_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = _get_or_add(borders, f"w:{border_name}")
        border.set(qn("w:val"), "nil")


def configure_cell(cell, width_dxa: int, *, borderless: bool = True) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    _set_dxa(_get_or_add(tc_pr, "w:tcW"), width_dxa)

    tc_mar = _get_or_add(tc_pr, "w:tcMar")
    for side, value in CELL_MARGIN_DXA.items():
        _set_dxa(_get_or_add(tc_mar, f"w:{side}"), value)
    if borderless:
        remove_cell_borders(cell)


def configure_table_geometry(
    table,
    widths: Sequence[int] = TABLE_COLUMN_WIDTHS_DXA,
    *,
    bordered: bool = False,
) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    _set_dxa(_get_or_add(tbl_pr, "w:tblW"), sum(widths))

    tbl_ind = _get_or_add(tbl_pr, "w:tblInd")
    _set_dxa(tbl_ind, 0)

    layout = _get_or_add(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    justification = _get_or_add(tbl_pr, "w:jc")
    justification.set(qn("w:val"), "left")

    tbl_borders = _get_or_add(tbl_pr, "w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = _get_or_add(tbl_borders, f"w:{border_name}")
        border.set(qn("w:val"), "single" if bordered else "nil")
        if bordered:
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "808080")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def add_key_value_table(doc: docx.Document, rows: Sequence[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=3)
    configure_table_geometry(table)

    emphasized_labels = {"Nama", "Nama Siswa", "NIP", "NIS / NISN"}
    for row, (label, value) in zip(table.rows, rows):
        for cell, width in zip(row.cells, TABLE_COLUMN_WIDTHS_DXA):
            configure_cell(cell, width)

        values = (label, ":", value)
        alignments = (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT)
        for index, (cell, text, alignment) in enumerate(zip(row.cells, values, alignments)):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(text)
            set_run_font(run, bold=(index == 2 and label in emphasized_labels))


def add_students_table(doc: docx.Document) -> None:
    widths = (700, 3300, 3100, 1271)
    table = doc.add_table(rows=4, cols=4)
    configure_table_geometry(table, widths, bordered=True)
    row_values = (
        ("No.", "Nama Siswa", "NIS / NISN", "Kelas"),
        ("{%tr for student in students %}", "", "", ""),
        ("{{ loop.index }}", "{{ student.nama }}", "{{ student.nis }} / {{ student.nisn }}", "{{ student.kelas }}"),
        ("{%tr endfor %}", "", "", ""),
    )
    for row_index, (row, values) in enumerate(zip(table.rows, row_values)):
        for cell, width, value in zip(row.cells, widths, values):
            configure_cell(cell, width, borderless=False)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index != 2 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, bold=(row_index == 0))


def add_signature(doc: docx.Document) -> None:
    signature_indent_cm = 9.0
    add_text_paragraph(
        doc,
        "Wonosari, {{ tanggal_surat }}",
        space_before_pt=8,
        keep_with_next=True,
        left_indent_cm=signature_indent_cm,
    )
    add_text_paragraph(
        doc,
        "{{ penandatangan_jabatan }}",
        keep_with_next=True,
        left_indent_cm=signature_indent_cm,
    )
    add_text_paragraph(
        doc,
        "\n\n",
        keep_with_next=True,
        left_indent_cm=signature_indent_cm,
    )
    add_text_paragraph(
        doc,
        "{{ penandatangan_nama }}",
        bold=True,
        underline=True,
        keep_with_next=True,
        left_indent_cm=signature_indent_cm,
    )

    # Tag {%p ... %} menghapus paragraf kontrolnya saat docxtpl merender.
    add_text_paragraph(
        doc,
        "{%p if penandatangan_id %}",
        left_indent_cm=signature_indent_cm,
    )
    add_text_paragraph(
        doc,
        "{{ penandatangan_id_label }} {{ penandatangan_id }}",
        left_indent_cm=signature_indent_cm,
    )
    add_text_paragraph(
        doc,
        "{%p endif %}",
        left_indent_cm=signature_indent_cm,
    )


def build_template(spec: TemplateSpec) -> Path:
    target = ACTIVE_TEMPLATE_DIR / spec.filename
    if target.resolve() == MASTER_DOC.resolve():
        raise RuntimeError("Master immutable tidak boleh menjadi target output.")

    doc = docx.Document(MASTER_DOC)
    retain_kop_only(doc)
    configure_default_font(doc)
    configure_page(doc)

    add_text_paragraph(
        doc,
        spec.title,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=TITLE_FONT_SIZE_PT,
        bold=True,
        underline=True,
        space_before_pt=10,
        keep_with_next=True,
    )
    add_text_paragraph(
        doc,
        "Nomor: {{ nomor_surat }}",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=8,
        keep_with_next=True,
    )
    add_text_paragraph(
        doc,
        spec.opening,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after_pt=6,
        keep_with_next=True,
    )
    if spec.multi_students:
        add_students_table(doc)
        add_text_paragraph(doc, "Rincian kegiatan:", space_before_pt=6, keep_with_next=True)
    add_key_value_table(doc, spec.rows)
    add_text_paragraph(
        doc,
        spec.closing,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_before_pt=8,
        space_after_pt=4,
    )
    add_signature(doc)

    doc.core_properties.author = "SMAN 2 Wonosari"
    doc.core_properties.last_modified_by = "E-Surat SMADA"
    doc.save(target)
    return target


def _variables_in(values: Iterable[str]) -> set[str]:
    variables: set[str] = set()
    for value in values:
        variables.update(re.findall(r"\{\{\s*([A-Za-z_][A-Za-z0-9_]*)", value))
    return variables


def expected_variables(spec: TemplateSpec) -> set[str]:
    source_values = [spec.title, spec.opening, spec.closing, "{{ nomor_surat }}"]
    source_values.extend(value for _, value in spec.rows)
    variables = _variables_in(source_values)
    if spec.multi_students:
        variables.add("students")
    variables.update(
        {
            "tanggal_surat",
            "penandatangan_jabatan",
            "penandatangan_nama",
            "penandatangan_id_label",
            "penandatangan_id",
        }
    )
    return variables


def audit_template(path: Path, spec: TemplateSpec) -> None:
    doc = docx.Document(path)
    expected_table_count = 2 if spec.multi_students else 1
    if len(doc.tables) != expected_table_count:
        raise RuntimeError(
            f"{path.name}: diharapkan {expected_table_count} tabel data, ditemukan {len(doc.tables)}."
        )
    detail_table = doc.tables[-1]
    if len(detail_table.rows) != len(spec.rows):
        raise RuntimeError(
            f"{path.name}: baris tabel {len(detail_table.rows)}, seharusnya {len(spec.rows)}."
        )

    with zipfile.ZipFile(path) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    actual_variables = set(DocxTemplate(str(path)).get_undeclared_template_variables())
    if actual_variables != expected_variables(spec):
        missing = sorted(expected_variables(spec) - actual_variables)
        unexpected = sorted(actual_variables - expected_variables(spec))
        raise RuntimeError(f"{path.name}: variabel hilang={missing}, tak dikenal={unexpected}.")

    drawing_count = document_xml.count("<w:drawing")
    if drawing_count != 1:
        raise RuntimeError(f"{path.name}: kop harus memiliki satu drawing; ditemukan {drawing_count}.")

    detail_xml = detail_table._tbl.xml
    tbl_width = re.search(r'<w:tblW\b[^>]*\bw:w="(\d+)"', detail_xml)
    grid_widths = tuple(
        int(value)
        for value in re.findall(r'<w:gridCol w:w="(\d+)"', detail_xml)[:3]
    )
    first_row = re.search(r"<w:tr(?:\s[^>]*)?>.*?</w:tr>", detail_xml, re.DOTALL)
    cell_widths = ()
    if first_row:
        cell_widths = tuple(
            int(value)
            for value in re.findall(r'<w:tcW\b[^>]*\bw:w="(\d+)"', first_row.group(0))[:3]
        )
    if not tbl_width or int(tbl_width.group(1)) != TABLE_WIDTH_DXA:
        raise RuntimeError(f"{path.name}: tblW tidak konsisten.")
    if grid_widths != TABLE_COLUMN_WIDTHS_DXA or cell_widths != TABLE_COLUMN_WIDTHS_DXA:
        raise RuntimeError(
            f"{path.name}: geometri tabel grid={grid_widths}, cell={cell_widths}."
        )
    if "<w:trHeight" in document_xml:
        raise RuntimeError(f"{path.name}: ditemukan fixed row height.")


def package_content_hash(path: Path) -> str:
    """Hash isi entry ZIP; timestamp container ZIP sengaja diabaikan."""
    digest = hashlib.sha256()
    with zipfile.ZipFile(path) as package:
        for name in sorted(package.namelist()):
            digest.update(name.encode("utf-8"))
            digest.update(b"\0")
            digest.update(package.read(name))
            digest.update(b"\0")
    return digest.hexdigest()


def main() -> None:
    if not MASTER_DOC.exists():
        raise FileNotFoundError(f"Master template tidak ditemukan: {MASTER_DOC}")
    if MASTER_DOC.name in {spec.filename for spec in TEMPLATE_SPECS}:
        raise RuntimeError("Master immutable tercantum sebagai output.")

    master_hash = package_content_hash(MASTER_DOC)
    for spec in TEMPLATE_SPECS:
        target = build_template(spec)
        audit_template(target, spec)
        print(f"[OK] {target.name}: {len(spec.rows)} baris, struktur dan placeholder valid.")

    if package_content_hash(MASTER_DOC) != master_hash:
        raise RuntimeError("Master immutable berubah selama build.")
    print("[OK] Master immutable tidak berubah.")


if __name__ == "__main__":
    main()
